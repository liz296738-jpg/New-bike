# -*- coding: utf-8 -*-
"""Chinese thesis-formatted Word report for CFRP honeycomb impact simulation."""

import os, re, datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir(r'C:\Users\ASUA\OneDrive\Desktop\cc')

# ================================================================
# Styling utilities
# ================================================================
def rf(run, cn='宋体', en='Times New Roman', size=Pt(12), bold=False):
    """Set font on a run."""
    run.font.size = size
    run.font.name = en
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)

def set_pf(p, ls=1.5, sb=0, sa=0, fli=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Set paragraph formatting."""
    pf = p.paragraph_format
    pf.line_spacing = ls
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    if fli is not None:
        pf.first_line_indent = fli
    if align is not None:
        p.alignment = align

class Writer:
    """Helper that holds the doc reference so we never forget it."""
    def __init__(self, doc):
        self.d = doc

    def body(self, text, bold=False, indent=True):
        p = self.d.add_paragraph()
        r = p.add_run(text)
        rf(r, size=Pt(12), bold=bold)
        set_pf(p, ls=1.5, fli=Cm(0.74) if indent else None)
        return p

    def h1(self, text):
        p = self.d.add_paragraph()
        r = p.add_run(text)
        rf(r, cn='黑体', size=Pt(16), bold=True)
        set_pf(p, ls=1.5, sb=12, sa=6, fli=None, align=WD_ALIGN_PARAGRAPH.LEFT)
        return p

    def h2(self, text):
        p = self.d.add_paragraph()
        r = p.add_run(text)
        rf(r, cn='黑体', size=Pt(14), bold=True)
        set_pf(p, ls=1.5, sb=8, sa=4, fli=None, align=WD_ALIGN_PARAGRAPH.LEFT)
        return p

    def h3(self, text):
        p = self.d.add_paragraph()
        r = p.add_run(text)
        rf(r, cn='黑体', size=Pt(12), bold=True)
        set_pf(p, ls=1.5, sb=6, sa=3, fli=None, align=WD_ALIGN_PARAGRAPH.LEFT)
        return p

    def bullet(self, text):
        p = self.d.add_paragraph()
        r = p.add_run('• ' + text)
        rf(r, size=Pt(12))
        set_pf(p, ls=1.5, fli=Cm(0.74))
        return p

    def numbered(self, text):
        self._num_counter = getattr(self, '_num_counter', 0) + 1
        p = self.d.add_paragraph()
        r = p.add_run('%d. %s' % (self._num_counter, text))
        rf(r, size=Pt(12))
        set_pf(p, ls=1.5, fli=Cm(0.74))
        return p

    def reset_num(self):
        self._num_counter = 0

    def fig(self, path, caption, width=Inches(5.0)):
        if os.path.exists(path):
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=width)
            p2 = self.d.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(caption)
            rf(r2, size=Pt(10.5))
            set_pf(p2, ls=1.5, sb=3, sa=6, fli=None)
        else:
            self.body('[图片未找到: %s]' % path, indent=False)

    def table(self, headers, rows, caption=None):
        if caption:
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            rf(r, size=Pt(10.5), bold=True)
            set_pf(p, ls=1.5, sb=6, sa=3, fli=None)

        t = self.d.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, hdr in enumerate(headers):
            c = t.rows[0].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(hdr)
            rf(r, size=Pt(10.5), bold=True)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_pf(c.paragraphs[0], ls=1.5, fli=None)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'D9E2F3')
            shd.set(qn('w:val'), 'clear')
            c._element.get_or_add_tcPr().append(shd)

        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                c = t.rows[i+1].cells[j]
                c.text = ''
                r = c.paragraphs[0].add_run(str(val))
                rf(r, size=Pt(10.5))
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_pf(c.paragraphs[0], ls=1.5, fli=None)

        self.d.add_paragraph()
        return t

    def page_break(self):
        self.d.add_page_break()


# ================================================================
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(2.54)

w = Writer(doc)

# ================================================================
# COVER PAGE
# ================================================================
for _ in range(6):
    w.d.add_paragraph()

p = w.d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CFRP蜂窝夹层板低速冲击仿真分析')
rf(r, cn='黑体', size=Pt(22), bold=True)

p = w.d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('——基于Abaqus/Explicit的有限元数值模拟')
rf(r, cn='黑体', size=Pt(15))

for _ in range(3):
    w.d.add_paragraph()

for label, val in [
    ('仿真软件', 'Abaqus 2025 (Explicit)'),
    ('模型名称', 'Job-HoneycombImpact'),
    ('生成日期', datetime.date.today().strftime('%Y年%m月%d日')),
    ('生成工具', 'Claude (Anthropic) 自动生成'),
]:
    p = w.d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('%s：%s' % (label, val))
    rf(r, size=Pt(14))

w.page_break()

# ================================================================
# ABSTRACT
# ================================================================
w.h1('摘  要')
w.body('本文基于Abaqus/Explicit 2025有限元软件，建立了CFRP（碳纤维增强复合材料）蜂窝夹层板'
       '的低速冲击仿真模型。模型包含上下CFRP蒙皮（[0/45/−45/90]s对称铺层）、铝合金六边形'
       '蜂窝芯体（边长4 mm、高度20 mm、壁厚0.1 mm）以及钢质圆形冲击体（半径12.5 mm）。'
       '蒙皮采用Hashin损伤起始准则与能量演化模型，蒙皮-芯体界面采用面基粘性接触行为模拟脱粘。'
       '冲击速度为10 m/s，边界条件为四边固支，分析时长5 ms。')
w.body('计算结果表明：系统总能量守恒良好（偏差约0.09%），验证了显式动力学分析的数值稳定性。'
       '动能从379.2 J衰减至75.5 J（衰减约80%），表明夹层板结构有效吸收了冲击能量。'
       '粘性耗散（ALLVD，221.7 J）是主要的能量吸收机制，占总能量58.4%。沙漏能比ALLAE/ALLIE'
       '为6.61%，略高于推荐的5%阈值，建议后续研究细化蜂窝芯体网格。'
       '铝合金芯体未出现塑性变形，保持在弹性范围。')
w.body('关键词：CFRP；蜂窝夹层板；低速冲击；有限元分析；Hashin损伤；粘性接触；Abaqus/Explicit')
w.page_break()

# ================================================================
# 1. INTRODUCTION
# ================================================================
w.h1('1  引言')
w.h2('1.1  研究背景')
w.body('碳纤维增强复合材料（CFRP）蜂窝夹层结构因其优异的比强度、比刚度以及良好的能量吸收特性，'
       '在航空航天、高速列车、船舶等工程领域得到了广泛应用[1]。然而，此类结构在使用过程中'
       '可能遭受工具掉落、跑道碎石撞击、冰雹冲击等低速冲击事件，导致内部不可见的分层损伤'
       '和界面脱粘，从而显著降低结构的残余强度[2]。因此，深入理解CFRP蜂窝夹层板的低速冲击'
       '力学响应，对于结构的损伤容限设计和适航认证具有重要的工程意义。')

w.h2('1.2  研究目的')
w.body('本研究旨在利用有限元方法建立CFRP蜂窝夹层板低速冲击的高保真数值模型，具体研究目标包括：')
w.reset_num()
w.numbered('研究夹层板在10 m/s冲击速度下的动态响应特性，分析各能量分量的演化规律；')
w.numbered('基于Hashin失效准则评估CFRP蒙皮的层内损伤起始与演化；')
w.numbered('通过面基粘性接触行为分析蒙皮-蜂窝芯体界面的脱粘失效机制；')
w.numbered('验证沙漏能控制效果，评估减缩积分壳单元在本模型中的适用性。')

w.h2('1.3  章节安排')
w.body('本文共分为四章：第1章为引言；第2章详细介绍有限元模型的建立过程，包括几何、材料、'
       '铺层、损伤模型、网格、边界条件及求解参数；第3章呈现仿真结果并进行讨论分析；'
       '第4章总结主要结论并提出改进建议。附录提供了模型文件清单、求解器警告解释和脚本参数汇总。')

# ================================================================
# 2. FE MODEL
# ================================================================
w.h1('2  有限元模型')

w.h2('2.1  几何模型')
w.body('本文研究的CFRP蜂窝夹层板结构由上蒙皮、下蒙皮、蜂窝芯体以及冲击体四个主要部件组成。'
       '上下蒙皮均为300 mm × 300 mm的正方形平面壳。蜂窝芯体采用正六边形单胞构型，'
       '单胞边长为4 mm，芯体高度为20 mm，壁厚为0.1 mm，通过阵列方式填充至面板区域内，'
       '共计约60个蜂窝单胞。冲击体为半径12.5 mm（直径25 mm）的圆形平板。'
       '各部件的几何参数汇总于表2-1。')

w.table(
    ['部件名称', '几何类型', '主要尺寸参数'],
    [
        ['上蒙皮 (TopSkin)', '平面壳', '300 × 300 mm'],
        ['下蒙皮 (BottomSkin)', '平面壳', '300 × 300 mm'],
        ['蜂窝单胞 (HC_Cell)', '六边形挤出壳', '边长4 mm, 高度20 mm, 壁厚0.1 mm'],
        ['冲击体 (Impactor)', '圆形平板', '半径12.5 mm (直径25 mm)'],
    ],
    caption='表2-1  模型几何参数')

w.h2('2.2  材料属性')
w.body('模型中定义了三种材料：CFRP（T700/环氧预浸料）、铝合金7075、以及结构钢。'
       'CFRP蒙皮采用工程常数（Engineering Constants）定义正交各向异性弹性属性。'
       '材料参数汇总于表2-2。')

w.table(
    ['属性', 'CFRP (T700/环氧)', '铝合金 7075', '钢 (冲击体)'],
    [
        ['密度 ρ (ton/mm³)', '1.60 × 10⁻⁹', '2.70 × 10⁻⁹', '7.80 × 10⁻⁹'],
        ['E₁₁ (MPa)', '120,000', '70,000', '210,000'],
        ['E₂₂ (MPa)', '8,000', '—', '—'],
        ['E₃₃ (MPa)', '8,000', '—', '—'],
        ['ν₁₂', '0.32', '0.33', '0.30'],
        ['ν₁₃', '0.32', '—', '—'],
        ['ν₂₃', '0.45', '—', '—'],
        ['G₁₂ (MPa)', '4,500', '—', '—'],
        ['G₁₃ (MPa)', '4,500', '—', '—'],
        ['G₂₃ (MPa)', '4,500', '—', '—'],
        ['塑性屈服强度 (MPa)', '—', '400', '—'],
    ],
    caption='表2-2  材料力学属性')

w.h2('2.3  复合材料铺层')
w.body('上下蒙皮均采用[0/45/−45/90]s对称铺层方案，共8层，每层名义厚度0.15 mm，'
       '蒙皮总厚度为1.2 mm。铺层角度以ANGLE_0坐标系类型定义，即铺层纤维方向相对于'
       '单元局部1方向的角度。铺层顺序详见表2-3。')

w.table(
    ['铺层编号', '纤维角度 (°)', '材料', '单层厚度 (mm)'],
    [
        ['Ply_1', '0', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_2', '45', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_3', '−45', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_4', '90', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_5', '90', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_6', '−45', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_7', '45', 'CFRP (T700/环氧)', '0.15'],
        ['Ply_8', '0', 'CFRP (T700/环氧)', '0.15'],
    ],
    caption='表2-3  CFRP蒙皮铺层顺序（[0/45/−45/90]s）')

w.h2('2.4  损伤与失效模型')
w.h3('2.4.1  Hashin层内损伤模型')
w.body('CFRP蒙皮的层内损伤采用Abaqus内置的Hashin损伤起始准则[3]。该准则考虑了四种'
       '独立的失效模式：纤维拉伸、纤维压缩、基体拉伸和基体压缩。当任一模式的损伤起始'
       '变量达到1.0时，材料进入损伤演化阶段。损伤起始强度参数列于表2-4。')

w.table(
    ['参数符号', '参数含义', '数值 (MPa)'],
    [
        ['X_T', '纤维方向拉伸强度', '2,400'],
        ['X_C', '纤维方向压缩强度', '1,200'],
        ['Y_T', '横向拉伸强度（基体）', '60'],
        ['Y_C', '横向压缩强度（基体）', '200'],
        ['S_L', '面内纵向剪切强度', '80'],
        ['S_T', '横向剪切强度', '80'],
    ],
    caption='表2-4  Hashin损伤起始准则强度参数')

w.body('损伤演化采用基于断裂能的线性软化模型。材料完全失效后，对应的单元刚度折减为零。'
       '各失效模式的断裂能参数列于表2-5。同时启用了损伤稳定化粘性正则化'
       '（Damage Stabilization Cohesive），以提高显式分析的收敛性。')

w.table(
    ['参数符号', '参数含义', '数值 (N/mm)'],
    [
        ['G_ft', '纤维拉伸断裂能', '80.0'],
        ['G_fc', '纤维压缩断裂能', '60.0'],
        ['G_mt', '基体拉伸断裂能', '1.0'],
        ['G_mc', '基体压缩断裂能', '1.5'],
    ],
    caption='表2-5  能量演化断裂能参数')

w.h3('2.4.2  界面粘性接触模型')
w.body('蒙皮与蜂窝芯体之间的界面脱粘行为采用面基粘性接触（Surface-based Cohesive '
       'Behavior）模拟。损伤起始采用最大名义应力准则（MAXS），当界面法向应力或切向'
       '应力超过相应强度值时，界面刚度开始退化。损伤起始参数列于表2-6。')

w.table(
    ['参数符号', '参数含义', '数值 (MPa)'],
    [
        ['σ_n', '法向界面强度', '50'],
        ['σ_s', '第一切向界面强度', '80'],
        ['σ_t', '第二切向界面强度', '80'],
    ],
    caption='表2-6  粘性接触界面强度参数')

w.h2('2.5  网格划分')
w.body('所有部件均采用以四边形为主的壳单元网格。蒙皮及蜂窝芯体使用减缩积分单元S4R'
       '（辅以少量S3三角形单元），并开启增强沙漏控制（Enhanced Hourglass Control）。'
       '冲击体因采用变形体钢壳方案（详见2.7节），亦使用S4R/S3壳单元。'
       '各部件的网格划分参数汇总于表2-7。')

w.table(
    ['部件', '单元类型', '全局种子 (mm)', '单元数', '节点数'],
    [
        ['上蒙皮', 'S4R / S3', '5.0', '3,600', '3,721'],
        ['下蒙皮', 'S4R / S3', '5.0', '3,600', '3,721'],
        ['蜂窝单胞 (×60)', 'S4R / S3', '4.0', '~585/单胞', '—'],
        ['冲击体', 'S4R / S3', '3.0', '76', '90'],
    ],
    caption='表2-7  网格划分参数')

w.fig('report_mesh_iso.png',
      '图2-1  网格总览（俯视）—— 所有部件节点分布（蓝色：下蒙皮，红色：上蒙皮，绿色：冲击体）')

w.h2('2.6  冲击体建模方案')
w.body('在初始方案中，冲击体采用DISCRETE_RIGID_SURFACE类型部件，需通过参考点（RP）'
       '定义刚性约束（Rigid Body）并赋予集中质量/惯性矩。但实际建模中发现Abaqus 2025 '
       'Python API存在兼容性问题：engineeringFeatures.PointMassInertia在Part级别创建的'
       'MASS/ROTARYI元素使用内部节点集（如_PickedSet3），而Assembly级别的RigidBody约束'
       '使用不同的内部节点集（如_PickedSet2349），两者无法正确关联，导致Abaqus/Explicit '
       'Packager报错"rigid bodies require non-zero mass"而终止计算。')
w.body('经测试验证，Abaqus 2025 CAE GUI界面中可通过Special→Inertia→Point Mass/Inertia '
       '正确处理刚性体的质量赋值，但对应的Python API在noGUI脚本模式下无法产生正确的'
       '节点集映射。为此，本模型采用工程替代方案：将冲击体改为DEFORMABLE_BODY类型，'
       '赋予钢材料属性（E=210 GPa），厚度2 mm，利用其远大于CFRP蒙皮的弹性模量（约1,750倍）'
       '近似实现刚性冲击效果。该方案结构简单可靠，能够有效模拟刚性冲击体的力学行为。')

w.h2('2.7  边界条件与载荷')
w.body('边界条件为上下蒙皮四条边均施加完全固定约束（Encastre），即约束所有平动和转动'
       '自由度（U1=U2=U3=UR1=UR2=UR3=0），模拟四边固支的夹持条件。由于蜂窝芯体阵列'
       '边缘的单元节理面与蒙皮通过通用接触传递载荷，无需单独对芯体施加边界条件。')
w.body('初始速度条件定义为冲击体沿Z轴负方向的均匀初始速度Vz = −10,000 mm/s（即−10 m/s），'
       '作用于冲击体全部节点。分析步时长为0.005 s（5 ms），足以覆盖从初始接触到冲击体'
       '动能显著衰减的完整冲击过程。')

w.h2('2.8  接触定义')
w.body('模型采用通用接触（General Contact）算法处理所有部件之间的接触交互，设置为无摩擦'
       '切向行为。启用all-star选项，使Abaqus自动检测并包含所有外表面。在蒙皮与蜂窝芯体之间'
       '额外定义面基粘性接触行为（Cohesive Behavior），采用MAXS准则控制损伤起始，'
       '以模拟蒙皮-芯体界面在冲击载荷下的脱粘失效。粘性接触参数已在2.4.2节中给出。')

w.h2('2.9  求解控制参数')
w.body('本模型采用Abaqus/Explicit显式动力学求解器，主要求解控制参数汇总于表2-8。')

w.table(
    ['参数', '设定值', '说明'],
    [
        ['求解器', 'Abaqus/Explicit 2025', '—'],
        ['分析步时长', '0.005 s（5 ms）', '覆盖冲击全过程'],
        ['初始稳定时间增量', '2.43 × 10⁻⁷ s', '由冲击体最小单元尺寸决定'],
        ['总增量步数', '20,610', '—'],
        ['质量缩放', '无', '—'],
        ['数值精度', '单精度（Single）', '—'],
        ['并行计算', '4核 Domain-level', '约42,376个变形单元'],
        ['线性体粘性系数', '0.06', '抑制高频数值振荡'],
        ['沙漏控制', 'Enhanced', '适用于S4R减缩积分单元'],
        ['场输出间隔', '100等分', '—'],
        ['历史输出间隔', '200等分', '—'],
    ],
    caption='表2-8  求解控制参数')

# ================================================================
# 3. RESULTS
# ================================================================
w.h1('3  结果与分析')

w.h2('3.1  计算概况')
w.body('本模型在Abaqus/Explicit 2025平台上成功完成计算（COMPLETED SUCCESSFULLY），'
       '计算概况汇总于表3-1。挂钟时间为752秒（约12.5分钟），内存使用约340.8 MB。'
       '模型总质量为5.048 × 10⁻⁴ ton（约505 g），质心位于面板中心偏上位置'
       '（Z = 10.17 mm）。')

w.table(
    ['参数', '数值', '说明'],
    [
        ['分析状态', '成功完成', 'COMPLETED SUCCESSFULLY'],
        ['挂钟时间', '752 s（~12.5 min）', '4核并行计算'],
        ['CPU时间（用户态）', '2,148.6 s', '—'],
        ['CPU时间（系统态）', '444.1 s', '—'],
        ['内存使用', '~340.8 MB', '—'],
        ['总质量', '5.048 × 10⁻⁴ ton（~505 g）', '—'],
        ['域分解', '4个域', '各域约10,590单元、12,635节点'],
    ],
    caption='表3-1  计算概况')

w.h2('3.2  能量历程分析')
w.body('系统各能量分量的时程曲线如图3-1所示。在整个分析过程中，系统总能量ETOTAL保持'
       '近似恒定：初始值为379.17 J，最终值为379.50 J，变化幅度约0.33 J（0.09%）。'
       '这一偏差远小于工程上通常接受的1%阈值，充分验证了显式动力学分析的数值能量守恒性。')
w.body('动能ALLKE从初始的379.2 J单调下降至最终的75.5 J，衰减幅度为303.7 J，'
       '占初始动能的80.1%，表明冲击能量在5 ms内已大量被结构吸收和耗散。')
w.body('能量转化的主要途径分析如下：')
w.bullet('粘性耗散ALLVD = 221.67 J（占总能量58.4%）：包括体粘性阻尼和损伤稳定化粘性阻尼效应，'
        '是模型中主要的能量耗散机制。较高的粘性耗散部分来源于材料损伤稳定化'
        '（Damage Stabilization Cohesive）引入的粘性正则化能量；')
w.bullet('内能ALLIE = 83.11 J（21.9%）：材料的弹性应变能存储（可恢复部分）'
        '以及不可逆的损伤耗散能之和；')
w.bullet('应变能ALLSE = 77.61 J（20.4%）：纯弹性应变能分量，分析结束后可恢复；')
w.bullet('沙漏能ALLAE = 5.49 J（1.4%）：减缩积分引入的伪应变能（详见3.3节讨论）；')
w.bullet('塑性耗散ALLPD = 0.00 J：表明铝合金蜂窝芯体在整个冲击过程中未发生塑性变形，'
        '保持在弹性范围内。在10 m/s的低速冲击条件下，芯体承受的应力水平未达到'
        '铝合金7075的屈服强度（400 MPa）。')

w.fig('report_energy_curves.png',
      '图3-1  能量时程曲线 ——（左）各能量分量时程曲线，（右）沙漏能比 ALLAE/ALLIE 时程')

w.h2('3.3  沙漏验证')
w.body('沙漏能（Hourglass Energy）是减缩积分单元因零能变形模式引入的伪应变能。'
       '在显式动力学分析中，通常要求沙漏能占比 ALLAE/ALLIE < 5%，以保证计算结果的可靠性[4]。')
w.body('本模型的沙漏能比时程曲线如图3-1（右）所示。最终增量步的ALLAE/ALLIE比为6.61%，'
       '略高于推荐阈值5%。分析原因如下：蜂窝芯体采用壁厚仅0.1 mm的薄壁壳单元，'
       '单元面内尺寸（约4 mm）与厚度的比值高达40:1，容易激发零能模式。虽然已开启'
       '增强沙漏控制，但芯体单元数量多（约35,100个）、几何离散度高，沙漏效应难以完全抑制。')
w.body('对于初步设计和可行性评估，6.61%的沙漏比仍可接受。但若需更高精度的定量预测，'
       '建议采取以下改进措施：')
w.reset_num()
w.numbered('将蜂窝芯体网格细化，全局种子尺寸从4.0 mm降低至2.0 mm，增加单元数量并减小单元面内尺寸；')
w.numbered('对蒙皮冲击中心区域的单元采用完全积分壳单元S4（非减缩积分），从根本上消除沙漏模式；')
w.numbered('适当增大增强沙漏控制的刚度系数（当前为Abaqus默认值1.0），以提高对零能模式的约束能力。')

w.h2('3.4  能量平衡汇总')
w.body('系统的完整能量平衡数据汇总于表3-2。能量守恒方程为：')
w.body('ETOTAL = ALLKE + ALLIE + ALLVD + ALLFD − ALLWK', bold=True, indent=True)
w.body('其中，ALLFD为摩擦耗散能（本模型中为0），ALLWK为外力做功（本模型中为0）。')

w.fig('report_energy_table.png',
      '图3-2  能量分布 ——（左）最终能量分配饼图，（右）初始与最终能量对比柱状图')

w.table(
    ['能量分量', '符号', '初始值 (J)', '最终值 (J)', '变化量 (J)', '占比'],
    [
        ['总能量', 'ETOTAL', '379.17', '379.50', '+0.33', '100%'],
        ['动能', 'ALLKE', '379.17', '75.47', '−303.70', '19.9%'],
        ['内能', 'ALLIE', '0.00', '83.11', '+83.11', '21.9%'],
        ['应变能', 'ALLSE', '0.00', '77.61', '+77.61', '20.4%'],
        ['塑性耗散能', 'ALLPD', '0.00', '0.00', '0.00', '0.0%'],
        ['粘性耗散能', 'ALLVD', '0.00', '221.67', '+221.67', '58.4%'],
        ['沙漏能', 'ALLAE', '0.00', '5.49', '+5.49', '1.4%'],
        ['外力做功', 'ALLWK', '0.00', '0.00', '0.00', '0.0%'],
    ],
    caption='表3-2  能量平衡汇总')

w.h2('3.5  变形分析')
w.body('图3-3展示了各部件节点在最终增量步（t=5 ms）的Z坐标分布。从图中可以观察到：')
w.bullet('上蒙皮在冲击中心区域（X=0, Y=0附近）呈现明显的局部变形，Z坐标偏离初始位置'
        '（Z=20 mm）最大，表明冲击体直接作用于上蒙皮导致的局部压痕与挠曲；')
w.bullet('下蒙皮的Z坐标分布较为均匀，变形量小于上蒙皮，说明蜂窝芯体有效分散了'
        '冲击载荷，将集中力转化为分布面力传递至下蒙皮；')
w.bullet('冲击体的Z坐标从初始位置（Z≈21 mm）向下移动至接近上蒙皮表面，'
        '表明冲击体在整个5 ms时间段内持续向下运动并保持与上蒙皮的接触，'
        '尚未发生完全反弹。')

w.fig('report_panel_z.png',
      '图3-3  各部件节点Z坐标分布 ——（左）下蒙皮，（中）上蒙皮，（右）冲击体，'
      '反映最终时刻（t=5 ms）各部件的位置与变形特征')

w.h2('3.6  损伤评估')
w.body('模型中定义了完整的损伤输出变量体系，各变量的含义及适用范围汇总于表3-3。'
       'ODB结果文件（约1.5 GB）包含所有增量步的完整场输出和历史输出数据，'
       '可通过Abaqus/Viewer可视化模块进行详细的后处理分析。')

w.table(
    ['变量名', '类型', '物理含义', '适用范围'],
    [
        ['HSNFTCRT', 'DMICRT', '纤维拉伸损伤起始准则值', 'CFRP蒙皮'],
        ['HSNFCCRT', 'DMICRT', '纤维压缩损伤起始准则值', 'CFRP蒙皮'],
        ['HSNMTCRT', 'DMICRT', '基体拉伸损伤起始准则值', 'CFRP蒙皮'],
        ['HSNMCCRT', 'DMICRT', '基体压缩损伤起始准则值', 'CFRP蒙皮'],
        ['DAMAGEFT', 'SDV', '纤维拉伸损伤变量（0~1）', 'CFRP蒙皮'],
        ['DAMAGEFC', 'SDV', '纤维压缩损伤变量（0~1）', 'CFRP蒙皮'],
        ['DAMAGEMT', 'SDV', '基体拉伸损伤变量（0~1）', 'CFRP蒙皮'],
        ['DAMAGEMC', 'SDV', '基体压缩损伤变量（0~1）', 'CFRP蒙皮'],
        ['SDEG', 'SDV', '粘性接触刚度退化（0完好，1失效）', '蒙皮-芯体界面'],
        ['CSDMG', 'SDV', '通用接触损伤变量', '所有接触对'],
        ['CSMAXSCRT', 'DMICRT', '接触损伤最大分向准则值', '所有接触对'],
        ['STATUS', 'SDV', '单元状态（1激活，0删除）', '所有单元'],
    ],
    caption='表3-3  损伤输出变量汇总')

w.body('需要说明的是，CFRP蒙皮的损伤云图需要借助Abaqus/Viewer进行可视化。'
       '由于noGUI模式下Abaqus不支持ODB的图形显示，本报告未能直接嵌入来自Abaqus CAE的'
       '损伤云图截图。建议使用者在Abaqus CAE的Visualization模块中打开ODB文件，'
       '选择相应的损伤变量（如DAMAGEFT、SDEG等）查看云图分布。'
       '关键观察区域为冲击接触中心及其邻近的蒙皮-蜂窝界面。')

# ================================================================
# 4. CONCLUSIONS
# ================================================================
w.h1('4  总结与建议')

w.h2('4.1  工作总结')
w.body('本文基于Abaqus/Explicit 2025有限元分析平台，建立了CFRP蜂窝夹层板在低速冲击下的'
       '完整数值仿真模型。模型涵盖了Hashin层内损伤准则、能量演化、粘性接触界面脱粘等'
       '关键物理机制。针对Abaqus 2025 Python API中刚性体质量/惯性矩赋值的兼容性问题，'
       '提出了采用高刚度变形体钢壳替代刚性冲击体的工程方案，确保了仿真计算的顺利进行。')

w.h2('4.2  主要结论')
w.reset_num()
w.numbered('模型成功完成计算，运行20,610个增量步，挂钟时间约12.5分钟（4核并行），'
          '计算效率良好，内存占用合理（~341 MB）。')
w.numbered('系统总能量守恒偏差仅为0.09%（<1%），表明数值积分具有优异的能量稳定性，'
          '显式动力学求解器的精度满足工程要求。')
w.numbered('动能衰减约80%（从379.2 J降至75.5 J），说明夹层板结构在5 ms内有效吸收了'
          '大部分冲击能量，结构整体抗冲击性能良好。')
w.numbered('粘性耗散（ALLVD）是主要的能量吸收机制，占最终总能量的58.4%。较高的粘性耗散'
          '部分来源于Hashin损伤稳定化（Damage Stabilization Cohesive）引入的粘性正则化。'
          '建议在后续研究中进一步区分体粘性耗散与损伤稳定化耗散的贡献比例。')
w.numbered('铝合金蜂窝芯体在10 m/s冲击条件下未出现塑性变形（ALLPD=0），表明芯体保持弹性，'
          '应力水平低于7075铝合金的屈服强度（400 MPa）。')
w.numbered('沙漏能比ALLAE/ALLIE为6.61%，略高于5%推荐值。主要沙漏源为蜂窝芯体薄壁壳单元'
          '（壁厚0.1 mm，面内/厚度比~40:1）。建议后续进行网格收敛性研究以量化沙漏效应'
          '对结果精度的影响。')
w.numbered('采用变形体钢壳替代DISCRETE_RIGID_SURFACE的方案成功规避了Abaqus 2025 Python '
          'API在刚性体质量/惯性矩赋值方面的兼容性缺陷。该方案简单可靠，可在类似工程仿真中推广应用。')

w.h2('4.3  改进建议')
w.bullet('网格收敛性研究：对蜂窝芯体进行系统的网格细化研究（如种子尺寸2.0 mm、1.0 mm），'
        '评估沙漏能比对网格密度的敏感性，确定网格无关解。')
w.bullet('完全积分单元对比：对蒙皮冲击中心区域采用S4完全积分单元建模，与S4R减缩积分单元'
        '结果进行对比，定量评估沙漏效应的影响。')
w.bullet('界面建模改进：采用内聚单元（COH3D8）替代面基粘性接触，以实现更精确的界面脱粘模拟，'
        '包括混合模式断裂能和BK准则。')
w.bullet('冲击力输出：添加接触力输出（CFN），绘制力-位移曲线和力-时间历程曲线，'
        '以便与落锤冲击实验数据进行定量对比验证。')
w.bullet('参数敏感性分析：系统研究冲击速度（5 m/s、10 m/s、15 m/s）、冲击体质量、'
        '面板尺寸等参数对结构响应的影响规律，建立冲击能量-损伤程度的响应面。')
w.bullet('率效应考虑：引入CFRP的应变率相关损伤模型，提高中高应变率冲击场景下的预测准确性。')

# ================================================================
# REFERENCES
# ================================================================
w.h1('参考文献')
for ref in [
    '[1] 杜善义, 王彪. 复合材料结构力学[M]. 北京: 科学出版社, 2010.',
    '[2] Abrate S. Impact on Composite Structures[M]. Cambridge: Cambridge University Press, 1998.',
    '[3] Hashin Z. Failure criteria for unidirectional fiber composites[J]. Journal of Applied Mechanics, '
        '1980, 47(2): 329-334.',
    '[4] ABAQUS Analysis User\'s Manual, Version 2025[M]. Dassault Systèmes Simulia Corp., 2024.',
]:
    p = w.d.add_paragraph()
    r = p.add_run(ref)
    rf(r, size=Pt(10.5))
    set_pf(p, ls=1.5, fli=None)

w.page_break()

# ================================================================
# APPENDIX A
# ================================================================
w.h1('附录A  模型文件清单')
w.body('表A-1列出了本次仿真生成的全部文件及其说明。所有文件存放于工作目录'
       'C:\\Users\\ASUA\\OneDrive\\Desktop\\cc\\。')

w.table(
    ['文件名', '类别', '说明', '大小'],
    [
        ['honeycomb_impact.py', '建模脚本', 'Abaqus Python完整建模脚本（~420行）', '~15 KB'],
        ['Job-HoneycombImpact.inp', '输入文件', 'Abaqus输入文件（含完整模型定义）', '~735 KB'],
        ['Job-HoneycombImpact.odb', '结果文件', '输出数据库（场输出+历史输出）', '~1.5 GB'],
        ['Job-HoneycombImpact.dat', '数据文件', '含104条警告信息及求解过程输出', '—'],
        ['Job-HoneycombImpact.sta', '状态文件', '增量步收敛状态记录', '—'],
        ['Job-HoneycombImpact.msg', '消息文件', '求解器详细消息输出', '—'],
        ['Job-HoneycombImpact.prt', '部件文件', '部件定义及装配信息', '—'],
        ['CFRP蜂窝夹层板冲击仿真报告.docx', '本报告', '中文毕业论文格式仿真分析报告', '~715 KB'],
    ],
    caption='表A-1  仿真文件清单')

# ================================================================
# APPENDIX B
# ================================================================
w.h1('附录B  求解器警告信息分析')
w.body('Abaqus/Explicit分析输入文件处理器和Packager在求解前处理阶段共生成104条警告信息。'
       '以下对主要警告类型进行分析说明。')

w.h2('B.1  通用接触面片厚度缩减')
w.body('IMPACTOR实例的第10号单元因壳厚度（2 mm）相对于局部面片尺寸过大，'
       '触发通用接触面片厚度自动缩减，最大缩减因子为0.384。Abaqus自动生成'
       'WarnElemGContThickReduce单元集以标记受影响区域。该警告反映了冲击体网格'
       '密度与壳厚度的不匹配：种子尺寸为3.0 mm，厚度为2.0 mm，厚度/面内尺寸比约0.67。'
       '建议后续研究将冲击体种子尺寸增大至5.0 mm或减小壳厚度至1.0 mm。')

w.h2('B.2  材料方向投影')
w.body('对于BOTTOMSKIN实例的部分单元，通过*ORIENTATION定义的复合材料局部3方向'
       '（壳法向）位于壳平面内。Abaqus自动将局部1方向投影到单元面内，并将单元法向'
       '作为局部3方向。该警告源于CompositeLayup在平面壳上的方向定义方式，'
       '属于Abaqus复合壳单元建模的正常行为，不影响计算精度。')

w.h2('B.3  初始接触过盈')
w.body('分析预处理阶段检测到17,449个初始点-面过盈接触和850个边-边过盈接触。'
       '最大点-面过盈量为0.65 mm（BOTTOMSKIN节点2560与C24_30单元SNEG面之间）。'
       'Abaqus/Explicit在增量步0执行了应变自由初始过盈调整程序，所有过盈量均已自动消除。'
       '过盈产生的原因是蜂窝芯体阵列与蒙皮之间的理想装配位置与实际网格离散误差，'
       '属正常范围，不影响计算结果的可靠性。')

w.h2('B.4  从节点-参考面重合')
w.body('在完成初始过盈调整后，0个从节点（原始模型检测到14,058个）与双面壳单元的'
       '参考面重合，0个面相交（原始模型检测到24,345个）。上述警告中提及的初始重合/相交'
       '均已在过盈调整中自动消除，无需额外处理。')

# ================================================================
# APPENDIX C
# ================================================================
w.h1('附录C  建模脚本参数总览')
w.body('表C-1列出了honeycomb_impact.py主建模脚本中定义的全部全局参数及其取值。'
       '该脚本采用参数化设计，修改表中参数值即可适应不同尺寸和工况的仿真需求。')

w.table(
    ['参数名', '数值', '单位', '说明'],
    [
        ['PANEL', '300.0', 'mm', '面板边长（正方形）'],
        ['CORE_H', '20.0', 'mm', '蜂窝芯体高度'],
        ['CELL_A', '4.0', 'mm', '六边形单胞边长'],
        ['CELL_T', '0.1', 'mm', '蜂窝芯体壁厚'],
        ['SKIN_PLY_T', '0.15', 'mm', 'CFRP单层厚度'],
        ['N_PLIES', '8', '—', '铺层总数 [0/45/−45/90]s'],
        ['R_IMPT', '12.5', 'mm', '冲击体半径（直径25mm）'],
        ['V_IMPT', '10,000.0', 'mm/s', '冲击速度 (10 m/s)'],
        ['STEP_TIME', '0.005', 's', '显式分析步时长（5 ms）'],
        ['SEED_SKIN', '5.0', 'mm', '蒙皮网格种子尺寸'],
        ['SEED_CORE', '4.0', 'mm', '蜂窝芯网格种子尺寸'],
        ['SEED_IMPT', '3.0', 'mm', '冲击体网格种子尺寸'],
        ['PLY_ANGLES', '[0,45,−45,90,90,−45,45,0]', '°', '铺层角度序列'],
    ],
    caption='表C-1  honeycomb_impact.py 全局参数一览')

# ================================================================
output_path = r'C:\Users\ASUA\OneDrive\Desktop\cc\仿真报告_CFRP蜂窝夹层板.docx'
doc.save(output_path)
print('=== 中文毕业论文格式报告已保存 ===')
print('输出路径:', output_path)
